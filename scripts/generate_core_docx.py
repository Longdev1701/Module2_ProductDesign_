import re
import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, hex_color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="C3C6D5", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), val)
        border.set(qn('w:sz'), sz)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        borders.append(border)
    insideV = OxmlElement('w:insideV')
    insideV.set(qn('w:val'), 'none')
    borders.append(insideV)
    tblPr.append(borders)

def main():
    md_path = r'e:\Projects\Project_ca_nhan\Module2\docs\mvp_core_usecases_deepdive.md'
    docx_path = r'e:\Projects\Project_ca_nhan\Module2\docs\Themis_LexiGuard_Dac_Ta_Usecase_MVP_Cot_Loi.docx'

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = docx.Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("PHÂN TÍCH CHUYÊN SÂU USE CASE CỐT LÕI (MVP CORE USE CASES)")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(16)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0, 50, 125)

    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Themis LexiGuard — Giải quyết Điểm đau Doanh nghiệp Xuất khẩu Nông sản")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(11)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(115, 119, 132)
    p_sub.paragraph_format.space_after = Pt(16)

    # Parse markdown headers and tables
    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue

        if line.startswith('# '):
            p = doc.add_paragraph()
            r = p.add_run(line[2:].strip())
            r.font.name = "Arial"
            r.font.size = Pt(16)
            r.font.bold = True
            r.font.color.rgb = RGBColor(0, 50, 125)
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.keep_with_next = True
            i += 1
            continue

        if line.startswith('## '):
            p = doc.add_paragraph()
            r = p.add_run(line[3:].strip())
            r.font.name = "Arial"
            r.font.size = Pt(13)
            r.font.bold = True
            r.font.color.rgb = RGBColor(0, 50, 125)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            i += 1
            continue

        if line.startswith('### '):
            p = doc.add_paragraph()
            r = p.add_run(line[4:].strip())
            r.font.name = "Arial"
            r.font.size = Pt(12)
            r.font.bold = True
            r.font.color.rgb = RGBColor(0, 50, 125)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            i += 1
            continue

        # Check if table block
        if line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1

            rows_data = []
            for t_line in table_lines:
                if '---|---' in t_line or '--- | ---' in t_line or '---|---|---' in t_line:
                    continue
                cells = [c.strip() for c in t_line.split('|')[1:-1]]
                if cells:
                    rows_data.append(cells)

            if rows_data:
                num_cols = max(len(r) for r in rows_data)
                table = doc.add_table(rows=len(rows_data), cols=num_cols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                set_table_borders(table, color="C3C6D5", sz="4")

                col_widths = [Inches(1.8), Inches(5.0)] if num_cols == 2 else [Inches(1.2), Inches(2.2), Inches(2.2), Inches(1.2)]

                for r_idx, row_cells in enumerate(rows_data):
                    row = table.rows[r_idx]
                    trPr = row._tr.get_or_add_trPr()
                    trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

                    for c_idx, cell_value in enumerate(row_cells):
                        if c_idx < len(row.cells):
                            cell = row.cells[c_idx]
                            cell.width = col_widths[c_idx] if c_idx < len(col_widths) else Inches(1.5)
                            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
                            p = cell.paragraphs[0]
                            p.paragraph_format.space_after = Pt(0)
                            p.paragraph_format.line_spacing = 1.15

                            if r_idx == 0:
                                set_cell_background(cell, "00327D")
                                run = p.add_run(re.sub(r'\*\*(.*?)\*\*', r'\1', cell_value))
                                run.font.name = "Arial"
                                run.font.size = Pt(9.5)
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(255, 255, 255)
                            else:
                                if c_idx == 0 and num_cols == 2:
                                    set_cell_background(cell, "F7F9FB")
                                else:
                                    set_cell_background(cell, "FFFFFF")

                                cell_val_clean = cell_value.replace('&nbsp;', ' ')
                                sub_lines = cell_val_clean.split('<br>')
                                for sl_idx, sl_text in enumerate(sub_lines):
                                    if sl_idx > 0:
                                        p = cell.add_paragraph()
                                        p.paragraph_format.space_after = Pt(0)
                                        p.paragraph_format.line_spacing = 1.15
                                    parts = re.split(r'(\*\*.*?\*\*)', sl_text)
                                    for part in parts:
                                        if not part:
                                            continue
                                        if part.startswith('**') and part.endswith('**'):
                                            r = p.add_run(part[2:-2])
                                            r.font.bold = True
                                        else:
                                            r = p.add_run(part)
                                        r.font.name = "Arial"
                                        r.font.size = Pt(9.5)
                                        r.font.color.rgb = RGBColor(25, 28, 30)

                p_space = doc.add_paragraph()
                p_space.paragraph_format.space_after = Pt(6)
            continue

        # Regular text paragraph
        if not line.startswith('```') and not line.startswith('---'):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if not part:
                    continue
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    r.font.bold = True
                else:
                    r = p.add_run(part)
                r.font.name = "Arial"
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(25, 28, 30)

        i += 1

    doc.save(docx_path)
    print(f"Successfully generated DOCX at {docx_path}")

if __name__ == '__main__':
    main()
