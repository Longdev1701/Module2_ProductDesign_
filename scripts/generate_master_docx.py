import re
import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def apply_font_to_run(run, font_name="Arial", size_pt=10, bold=False, italic=False, color_rgb=RGBColor(25, 28, 30)):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color_rgb

    # Explicitly append w:rFonts to ensure MS Word renders Vietnamese UTF-8 glyphs properly
    rPr = run._r.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}" w:cs="{font_name}"/>')
    rPr.append(rFonts)

def set_cell_background(cell, hex_color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="C3C6D5", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), val)
        border.set(qn('w:sz'), sz)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        borders.append(border)
    insideV = OxmlElement('w:insideV')
    insideV.set(qn('w:val'), 'none')
    borders.append(insideV)
    tblPr.append(borders)

def add_formatted_text(paragraph, text, base_size=10, base_color=RGBColor(25, 28, 30)):
    # Parse markdown code `...` and bold **...** and italic *...*
    parts = re.split(r'(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            apply_font_to_run(run, font_name="Consolas", size_pt=base_size-0.5, color_rgb=RGBColor(0, 50, 125))
        elif part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            apply_font_to_run(run, font_name="Arial", size_pt=base_size, bold=True, color_rgb=base_color)
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            apply_font_to_run(run, font_name="Arial", size_pt=base_size, italic=True, color_rgb=base_color)
        else:
            run = paragraph.add_run(part)
            apply_font_to_run(run, font_name="Arial", size_pt=base_size, color_rgb=base_color)

def main():
    md_path = r'e:\Projects\Project_ca_nhan\Module2\docs\THEMIS_LEXIGUARD_MASTER_SPECIFICATION.md'
    docx_path = r'e:\Projects\Project_ca_nhan\Module2\docs\Themis_LexiGuard_Tai_Lieu_Quy_Chuan_Doanh_Nghiep.docx'

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = docx.Document()

    # Normal Style default font
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(10)
    normal_style.font.color.rgb = RGBColor(25, 28, 30)

    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    lines = content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i].rstrip()

        if not line.strip():
            i += 1
            continue

        # Document Title (# )
        if line.startswith('# '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line[2:].strip())
            apply_font_to_run(run, font_name="Arial", size_pt=16, bold=True, color_rgb=RGBColor(0, 35, 111))
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # Subtitle (## )
        if line.startswith('## '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line[3:].strip())
            apply_font_to_run(run, font_name="Arial", size_pt=12, bold=True, color_rgb=RGBColor(0, 71, 171))
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(14)
            i += 1
            continue

        # Chapter Header (### )
        if line.startswith('### '):
            p = doc.add_paragraph()
            run = p.add_run(line[4:].strip())
            apply_font_to_run(run, font_name="Arial", size_pt=13, bold=True, color_rgb=RGBColor(0, 35, 111))
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            i += 1
            continue

        # Section Header (#### )
        if line.startswith('#### '):
            p = doc.add_paragraph()
            run = p.add_run(line[5:].strip())
            apply_font_to_run(run, font_name="Arial", size_pt=11, bold=True, color_rgb=RGBColor(0, 71, 171))
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            i += 1
            continue

        # Code Block (```)
        if line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1

            p_code = doc.add_paragraph()
            p_code.paragraph_format.space_before = Pt(6)
            p_code.paragraph_format.space_after = Pt(6)
            p_code.paragraph_format.line_spacing = 1.1

            table_code = doc.add_table(rows=1, cols=1)
            table_code.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell_code = table_code.rows[0].cells[0]
            cell_code.width = Inches(6.8)
            set_cell_background(cell_code, "F7F9FB")
            set_cell_margins(cell_code, top=140, bottom=140, left=180, right=180)

            p_c = cell_code.paragraphs[0]
            p_c.paragraph_format.space_after = Pt(0)
            p_c.paragraph_format.line_spacing = 1.15

            code_text = "\n".join(code_lines)
            run_c = p_c.add_run(code_text)
            apply_font_to_run(run_c, font_name="Consolas", size_pt=9, color_rgb=RGBColor(0, 50, 125))
            continue

        # Table block (|)
        if line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1

            rows_data = []
            for t_line in table_lines:
                if '---|---' in t_line or '--- | ---' in t_line or '---|---|---|---' in t_line:
                    continue
                cells = [c.strip() for c in t_line.split('|')[1:-1]]
                if cells:
                    rows_data.append(cells)

            if rows_data:
                num_cols = max(len(r) for r in rows_data)
                table = doc.add_table(rows=len(rows_data), cols=num_cols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                set_table_borders(table, color="C3C6D5", sz="4")

                col_widths = [Inches(1.8), Inches(1.5), Inches(1.8), Inches(1.7)]

                for r_idx, row_cells in enumerate(rows_data):
                    row = table.rows[r_idx]
                    trPr = row._tr.get_or_add_trPr()
                    trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

                    for c_idx, cell_value in enumerate(row_cells):
                        if c_idx < len(row.cells):
                            cell = row.cells[c_idx]
                            cell.width = col_widths[c_idx] if c_idx < len(col_widths) else Inches(1.5)
                            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
                            p = cell.paragraphs[0]
                            p.paragraph_format.space_after = Pt(0)
                            p.paragraph_format.line_spacing = 1.15

                            if r_idx == 0:
                                set_cell_background(cell, "00236F")
                                run = p.add_run(cell_value.replace('**', ''))
                                apply_font_to_run(run, font_name="Arial", size_pt=9.5, bold=True, color_rgb=RGBColor(255, 255, 255))
                            else:
                                if c_idx == 0:
                                    set_cell_background(cell, "F7F9FB")
                                else:
                                    set_cell_background(cell, "FFFFFF")
                                add_formatted_text(p, cell_value, base_size=9)

                p_space = doc.add_paragraph()
                p_space.paragraph_format.space_after = Pt(4)
            continue

        # Regular text or List item
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15

        if line.startswith('- ') or line.startswith('* '):
            p.paragraph_format.left_indent = Inches(0.25)
            add_formatted_text(p, "• " + line[2:].strip())
        elif re.match(r'^\d+\.\s', line):
            p.paragraph_format.left_indent = Inches(0.25)
            add_formatted_text(p, line.strip())
        else:
            add_formatted_text(p, line.strip())

        i += 1

    doc.save(docx_path)
    print(f"Successfully generated 100% Vietnamese Master Specification DOCX at {docx_path}")

if __name__ == '__main__':
    main()
