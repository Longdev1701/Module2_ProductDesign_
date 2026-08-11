import re
import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, hex_color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="CCCCCC", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), val)
        border.set(qn('w:sz'), sz)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        borders.append(border)
    insideV = OxmlElement('w:insideV')
    insideV.set(qn('w:val'), 'none')
    borders.append(insideV)
    tblPr.append(borders)

def main():
    md_path = r'e:\Projects\Project_ca_nhan\Module2\docs\usecases_doc_format.md'
    docx_path = r'e:\Projects\Project_ca_nhan\Module2\docs\Themis_LexiGuard_Dac_Ta_Usecase.docx'

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = docx.Document()

    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("BẢNG ĐẶC TẢ USE CASE HỆ THỐNG THEMIS LEXIGUARD")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(18)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0, 50, 125) # Deep navy #00327D

    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("AI Compliance Navigator for Agricultural Export — Enterprise Specifications")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(11)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(115, 119, 132)
    p_sub.paragraph_format.space_after = Pt(20)

    # Parse Sections (### 1.6.X...)
    uc_blocks = re.split(r'\n(?=### 1\.6\.\d+\.)', content)

    for block in uc_blocks:
        block = block.strip()
        if not block:
            continue

        # Header match
        header_match = re.match(r'### (1\.6\.\d+\. Usecase [^\n]+)', block)
        if not header_match:
            continue

        uc_title = header_match.group(1).strip()

        # Add section title
        p_uc = doc.add_paragraph()
        run_uc = p_uc.add_run(uc_title)
        run_uc.font.name = "Arial"
        run_uc.font.size = Pt(14)
        run_uc.font.bold = True
        run_uc.font.color.rgb = RGBColor(0, 50, 125)
        p_uc.paragraph_format.space_before = Pt(16)
        p_uc.paragraph_format.space_after = Pt(6)
        p_uc.paragraph_format.keep_with_next = True

        # Parse table rows
        rows_data = []
        table_lines = [line.strip() for line in block.split('\n') if line.strip().startswith('|')]

        for line in table_lines:
            if '---|---' in line or '--- | ---' in line:
                continue
            cells = [c.strip() for c in line.split('|')[1:-1]]
            if len(cells) >= 2:
                rows_data.append((cells[0], cells[1]))

        if not rows_data:
            continue

        # Create Word Table
        table = doc.add_table(rows=len(rows_data), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        set_table_borders(table, color="C3C6D5", sz="4")

        # Column widths
        col_widths = [Inches(1.8), Inches(5.0)]

        for row_idx, (attr, val) in enumerate(rows_data):
            row = table.rows[row_idx]

            # Prevent row split across pages
            trPr = row._tr.get_or_add_trPr()
            trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

            # Left Cell (Attribute)
            cell_left = row.cells[0]
            cell_left.width = col_widths[0]
            set_cell_margins(cell_left, top=120, bottom=120, left=150, right=150)
            p_left = cell_left.paragraphs[0]
            p_left.paragraph_format.space_after = Pt(0)
            p_left.paragraph_format.line_spacing = 1.15

            # Clean markdown bold ** text
            attr_clean = re.sub(r'\*\*(.*?)\*\*', r'\1', attr)
            run_left = p_left.add_run(attr_clean)
            run_left.font.name = "Arial"
            run_left.font.size = Pt(9.5)
            run_left.font.bold = True

            if row_idx == 0:
                # Header row styling
                set_cell_background(cell_left, "00327D")
                run_left.font.color.rgb = RGBColor(255, 255, 255)
            else:
                set_cell_background(cell_left, "F7F9FB")
                run_left.font.color.rgb = RGBColor(25, 28, 30)

            # Right Cell (Value / Description)
            cell_right = row.cells[1]
            cell_right.width = col_widths[1]
            set_cell_margins(cell_right, top=120, bottom=120, left=150, right=150)
            p_right = cell_right.paragraphs[0]
            p_right.paragraph_format.space_after = Pt(0)
            p_right.paragraph_format.line_spacing = 1.15

            val_clean = val.replace('&nbsp;', ' ')
            lines = val_clean.split('<br>')

            if row_idx == 0:
                set_cell_background(cell_right, "00327D")
                run_right = p_right.add_run(re.sub(r'\*\*(.*?)\*\*', r'\1', lines[0]))
                run_right.font.name = "Arial"
                run_right.font.size = Pt(9.5)
                run_right.font.bold = True
                run_right.font.color.rgb = RGBColor(255, 255, 255)
            else:
                set_cell_background(cell_right, "FFFFFF")
                for l_idx, l_text in enumerate(lines):
                    if l_idx > 0:
                        p_right = cell_right.add_paragraph()
                        p_right.paragraph_format.space_after = Pt(0)
                        p_right.paragraph_format.line_spacing = 1.15

                    # Parse bold parts
                    parts = re.split(r'(\*\*.*?\*\*)', l_text)
                    for part in parts:
                        if not part:
                            continue
                        if part.startswith('**') and part.endswith('**'):
                            r = p_right.add_run(part[2:-2])
                            r.font.bold = True
                        else:
                            r = p_right.add_run(part)
                        r.font.name = "Arial"
                        r.font.size = Pt(9.5)
                        r.font.color.rgb = RGBColor(25, 28, 30)

        # Add spacing after table
        p_space = doc.add_paragraph()
        p_space.paragraph_format.space_before = Pt(0)
        p_space.paragraph_format.space_after = Pt(8)

    doc.save(docx_path)
    print(f"Successfully generated DOCX at {docx_path}")

if __name__ == '__main__':
    main()
